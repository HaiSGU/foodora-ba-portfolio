from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class TableBuffer:
    header: list[str] | None = None
    rows: list[list[str]] | None = None

    def __post_init__(self) -> None:
        if self.rows is None:
            self.rows = []

    def add_row(self, row: list[str]) -> None:
        assert self.rows is not None
        self.rows.append(row)

    def is_empty(self) -> bool:
        return (self.header is None) and (not self.rows)


def _clean_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _is_table_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if "|" not in s:
        return False
    # ignore code fences or markdown separators outside tables
    return True


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    # Typical: |---|---:|---|
    if "|" not in s:
        return False
    parts = [p.strip() for p in s.strip("|").split("|")]
    if not parts:
        return False
    return all(re.fullmatch(r":?-{3,}:?", p) for p in parts)


def _parse_table_row(line: str) -> list[str]:
    parts = [p for p in line.strip().strip("|").split("|")]
    return [_clean_cell(p) for p in parts]


def _flush_table(doc: Document, buf: TableBuffer) -> None:
    if buf.is_empty():
        return

    header = buf.header
    rows = buf.rows or []
    if header is None:
        # treat first row as header if not set
        if not rows:
            return
        header = rows[0]
        rows = rows[1:]

    col_count = max(1, len(header))
    for r in rows:
        col_count = max(col_count, len(r))

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # header
    for i in range(col_count):
        cell_text = header[i] if i < len(header) else ""
        table.cell(0, i).text = cell_text

    # body
    for ri, row in enumerate(rows, start=1):
        for ci in range(col_count):
            cell_text = row[ci] if ci < len(row) else ""
            table.cell(ri, ci).text = cell_text

    # spacing paragraph
    doc.add_paragraph("")


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Title from first H1 if present
    title_done = False

    table_buf = TableBuffer()
    in_table = False

    for raw in lines:
        line = raw.rstrip("\n")

        # Skip horizontal rules
        if line.strip() in {"---", "***"}:
            if in_table:
                _flush_table(doc, table_buf)
                table_buf = TableBuffer()
                in_table = False
            continue

        # Table handling
        if _is_table_line(line):
            if _is_table_separator(line):
                # separator line: ignore
                in_table = True
                continue

            row = _parse_table_row(line)
            if not in_table:
                # starting a new table: flush any previous
                _flush_table(doc, table_buf)
                table_buf = TableBuffer()
                in_table = True

            if table_buf.header is None:
                table_buf.header = row
            else:
                table_buf.add_row(row)
            continue

        # leaving table
        if in_table:
            _flush_table(doc, table_buf)
            table_buf = TableBuffer()
            in_table = False

        s = line.strip()
        if not s:
            doc.add_paragraph("")
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()

            if level == 1 and not title_done:
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_done = True
                doc.add_paragraph("")
            else:
                heading_level = min(4, level)  # docx only has Heading 1..9; keep it compact
                doc.add_heading(content, level=heading_level)
            continue

        # blockquotes
        if s.startswith(">"):
            content = s.lstrip(">").strip()
            p = doc.add_paragraph(content)
            p.style = "Intense Quote" if "Intense Quote" in [st.name for st in doc.styles] else p.style
            continue

        # unordered list
        if s.startswith("- ") or s.startswith("* "):
            content = s[2:].strip()
            doc.add_paragraph(content, style="List Bullet")
            continue

        # ordered list (simple)
        m = re.match(r"^\d+\.\s+(.*)$", s)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            continue

        # bold lines like **Project:** ... (keep as normal paragraph; strip markdown bold markers)
        s2 = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
        doc.add_paragraph(s2)

    if in_table:
        _flush_table(doc, table_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        return 2

    md_path = Path(argv[1]).resolve()
    docx_path = Path(argv[2]).resolve()

    if not md_path.exists():
        print(f"Input not found: {md_path}")
        return 1

    convert_md_to_docx(md_path, docx_path)
    print(f"Wrote: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
